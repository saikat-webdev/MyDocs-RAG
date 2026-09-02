from pathlib import Path
from typing import List, Dict, Any, Tuple
import pypdf
import docx
from app.core.logging import logger

class DocumentParseError(Exception):
    pass

class ParserService:
    @staticmethod
    def parse_document(file_path: Path, file_type: str) -> Tuple[List[Dict[str, Any]], int]:
        """
        Parses document and returns:
        (pages_data, total_pages)
        where pages_data is a list of dicts: {"page_number": int, "text": str}
        """
        ext = file_type.lower().lstrip(".")
        if ext == "pdf":
            return ParserService.parse_pdf(file_path)
        elif ext == "docx":
            return ParserService.parse_docx(file_path)
        elif ext == "txt":
            return ParserService.parse_txt(file_path)
        elif ext == "md":
            return ParserService.parse_markdown(file_path)
        else:
            raise DocumentParseError(f"Unsupported document format: {file_type}")

    @staticmethod
    def parse_pdf(file_path: Path) -> Tuple[List[Dict[str, Any]], int]:
        pages = []
        try:
            reader = pypdf.PdfReader(str(file_path))
            total_pages = len(reader.pages)
            if total_pages == 0:
                raise DocumentParseError("PDF document contains no pages.")
            
            for idx, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                cleaned = page_text.strip()
                if cleaned:
                    pages.append({"page_number": idx, "text": cleaned})
            
            if not pages:
                raise DocumentParseError("Could not extract any readable text from the PDF.")
            
            return pages, total_pages
        except Exception as e:
            if isinstance(e, DocumentParseError):
                raise
            logger.error(f"Failed to parse PDF {file_path}: {e}", exc_info=True)
            raise DocumentParseError(f"PDF extraction failed: {str(e)}")

    @staticmethod
    def parse_docx(file_path: Path) -> Tuple[List[Dict[str, Any]], int]:
        try:
            doc = docx.Document(str(file_path))
            paragraphs = []
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    paragraphs.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            full_text = "\n\n".join(paragraphs).strip()
            if not full_text:
                raise DocumentParseError("DOCX document contains no text.")
            
            # DOCX doesn't have explicit page breaks naturally; treat as 1 logical section/page
            return [{"page_number": 1, "text": full_text}], 1
        except Exception as e:
            if isinstance(e, DocumentParseError):
                raise
            logger.error(f"Failed to parse DOCX {file_path}: {e}", exc_info=True)
            raise DocumentParseError(f"DOCX extraction failed: {str(e)}")

    @staticmethod
    def parse_txt(file_path: Path) -> Tuple[List[Dict[str, Any]], int]:
        try:
            content = None
            for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise DocumentParseError("Unable to decode text file with supported encodings.")
            
            cleaned = content.strip()
            if not cleaned:
                raise DocumentParseError("Text file is empty.")
            
            return [{"page_number": 1, "text": cleaned}], 1
        except Exception as e:
            if isinstance(e, DocumentParseError):
                raise
            logger.error(f"Failed to parse TXT {file_path}: {e}", exc_info=True)
            raise DocumentParseError(f"Text file extraction failed: {str(e)}")

    @staticmethod
    def parse_markdown(file_path: Path) -> Tuple[List[Dict[str, Any]], int]:
        try:
            pages, _ = ParserService.parse_txt(file_path)
            return pages, 1
        except Exception as e:
            raise DocumentParseError(f"Markdown extraction failed: {str(e)}")
